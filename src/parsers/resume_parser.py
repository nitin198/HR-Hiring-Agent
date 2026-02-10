"""Resume parser for extracting text from various file formats."""

import io
import zipfile
from typing import Iterable
from io import BytesIO
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

from docx import Document
from pypdf import PdfReader


class ResumeParser:
    """Parser for extracting text from resume files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md"}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

    @classmethod
    def parse(cls, file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """
        Parse resume file and extract text.

        Args:
            file_path: Path to the resume file
            file_content: Optional file content bytes (useful for uploaded files)

        Returns:
            Extracted text from the resume

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist and no content provided
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {extension}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        # Use provided content if available, otherwise read from file
        if file_content is not None:
            content_bytes = file_content
        else:
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            content_bytes = path.read_bytes()

        # Validate file size
        if len(content_bytes) > cls.MAX_FILE_SIZE:
            raise ValueError(
                f"File size ({len(content_bytes)} bytes) exceeds maximum allowed size "
                f"({cls.MAX_FILE_SIZE} bytes)"
            )

        # Parse based on file type
        if extension == ".pdf":
            return cls._parse_pdf(content_bytes)
        elif extension == ".doc":
            return cls._parse_doc(content_bytes)
        elif extension == ".docx":
            return cls._parse_docx(content_bytes)
        elif extension in {".txt", ".md"}:
            return content_bytes.decode("utf-8", errors="ignore")

        return ""

    @staticmethod
    def _parse_pdf(content_bytes: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(content_bytes)
            reader = PdfReader(pdf_file)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    def _parse_docx(content_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc_file = io.BytesIO(content_bytes)
            doc = Document(doc_file)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            # Fallback: extract all text nodes directly from the DOCX XML,
            # which can include headers, footers, and text boxes.
            try:
                xml_text = cls._extract_docx_xml_text(content_bytes)
                if xml_text:
                    text_parts.append(xml_text)
            except Exception:
                pass

            # OCR images embedded in DOCX (e.g., scanned resumes or name in header image)
            try:
                image_text = cls._ocr_docx_images(content_bytes)
                if image_text:
                    text_parts.append(image_text)
            except Exception:
                pass

            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def _extract_docx_xml_text(content_bytes: bytes) -> str:
        """Extract all text nodes from DOCX XML files."""
        text_parts = []
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as docx_zip:
            xml_files = [
                name
                for name in docx_zip.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]
            for name in xml_files:
                try:
                    xml_data = docx_zip.read(name)
                    root = ET.fromstring(xml_data)
                except Exception:
                    continue
                for elem in root.iter():
                    if elem.tag.endswith("}t") and elem.text:
                        text_parts.append(elem.text)
        return "\n".join(text_parts)

    @staticmethod
    def _ocr_docx_images(content_bytes: bytes) -> str:
        """Run OCR on images embedded in a DOCX file."""
        image_bytes_list = []
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as docx_zip:
            for name in docx_zip.namelist():
                if name.startswith("word/media/"):
                    image_bytes_list.append(docx_zip.read(name))
        return ResumeParser._ocr_images(image_bytes_list)

    @staticmethod
    def _ocr_images(images: Iterable[bytes]) -> str:
        """Extract text from a list of image byte payloads using OCR."""
        try:
            from PIL import Image
            import pytesseract
        except Exception:
            return ""

        text_parts = []
        for img_bytes in images:
            try:
                img = Image.open(BytesIO(img_bytes))
                text = pytesseract.image_to_string(img)
                if text and text.strip():
                    text_parts.append(text.strip())
            except Exception:
                continue
        return "\n".join(text_parts)

    @staticmethod
    def _parse_doc(content_bytes: bytes) -> str:
        """Extract text from DOC file."""
        try:
            try:
                import textract  # type: ignore
            except Exception as exc:
                raise ValueError("DOC parsing requires the textract library") from exc
            import tempfile

            with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as temp_file:
                temp_file.write(content_bytes)
                temp_file.flush()
                text = textract.process(temp_file.name)
                return text.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"Failed to parse DOC: {e}")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        lines = [line.strip() for line in text.split("\n")]
        lines = [line for line in lines if line]  # Remove empty lines

        # Join with single newlines
        cleaned = "\n".join(lines)

        # Remove multiple consecutive spaces
        import re
        cleaned = re.sub(r" {2,}", " ", cleaned)

        return cleaned.strip()

    @classmethod
    def parse_and_clean(cls, file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """
        Parse resume file and return cleaned text.

        Args:
            file_path: Path to the resume file
            file_content: Optional file content bytes

        Returns:
            Cleaned extracted text
        """
        raw_text = cls.parse(file_path, file_content)
        return cls.clean_text(raw_text)
